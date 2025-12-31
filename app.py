"""Flask backend for the TrustDoc.ai landing and document analysis MVP."""
from threads_review.routes import threads_review_bp
from flask import redirect

import os
from io import BytesIO
from typing import Dict, Tuple

import PyPDF2
import docx
import langdetect
import markdown
from dotenv import load_dotenv
from flask import Flask, jsonify, render_template, request
from flask_cors import CORS
from langdetect.lang_detect_exception import LangDetectException
from openai import OpenAI

load_dotenv()

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
OPENAI_MODEL = "gpt-4o-mini"

SYSTEM_PROMPT = (
    "You are a Kazakhstani legal expert working for TrustDoc.ai — an AI system for document and contract risk analysis. "
    "Analyze uploaded documents in the context of the laws and Constitution of the Republic of Kazakhstan, "
    "using the official database Adilet.gov.kz for legal reasoning when appropriate. "
    "Detect potential legal or contractual risks and cite relevant Kazakhstani legal norms briefly. "
    "Identify the document’s language (Kazakh, Russian, or English) automatically and respond only in that same language. "
    "Never mix multiple languages in one response. "
    "Keep your output concise, factual, and formatted in Markdown. "
    "If there are no significant legal risks, respond briefly: "
    "'✅ No significant legal risks under the laws of the Republic of Kazakhstan.' "
    "If risks exist, structure your output as follows:"
    "\n\n⚠️ Potential Legal Risks (based on Kazakh law):"
    "\n- [Short description of the issue]"
    "\n- [Relevant reference, e.g. 'May contradict Article 26 of the Constitution of the Republic of Kazakhstan' or 'Violates Article 375 of the Civil Code']"
    "\n\nKeep your answer under 1000 tokens."
)

TITLE_COPY: Dict[str, Tuple[str, str, str]] = {
    "index": (
        "TrustDoc.ai — Smart Document Analyzer",
        "TrustDoc.ai — Умный анализатор документов",
        "TrustDoc.ai — Құжаттарды ақылды талдаушы",
    ),
    "product": (
        "TrustDoc.ai — Smart Document Analyzer",
        "TrustDoc.ai — Умный анализатор документов",
        "TrustDoc.ai — Құжаттарды ақылды талдаушы",
    ),
    "about": (
        "About TrustDoc.ai",
        "О TrustDoc.ai",
        "TrustDoc.ai туралы",
    ),
    "project": (
        "TrustDoc.ai Project Vision",
        "Проектное видение TrustDoc.ai",
        "TrustDoc.ai жобалық көзқарасы",
    ),
    "news": (
        "TrustDoc.ai News",
        "Новости TrustDoc.ai",
        "TrustDoc.ai жаңалықтары",
    ),
    "gallery": (
        "TrustDoc.ai Gallery",
        "Галерея TrustDoc.ai",
        "TrustDoc.ai галереясы",
    ),
    "contacts": (
        "Contact TrustDoc.ai",
        "Контакты TrustDoc.ai",
        "TrustDoc.ai байланыс",
    ),
    "freelancers": (
        "TrustDoc.ai for Freelancers",
        "TrustDoc.ai для фрилансеров",
        "TrustDoc.ai фрилансерлерге",
    ),
    "startups": (
        "TrustDoc.ai for Startups",
        "TrustDoc.ai для стартапов",
        "TrustDoc.ai стартаптарға",
    ),
    "news_detail": (
        "TrustDoc.ai — Update",
        "TrustDoc.ai — Обновление",
        "TrustDoc.ai — Жаңалық",
    ),
}

DESCRIPTION_COPY: Dict[str, Tuple[str, str, str]] = {
    "index": (
        "Upload your contract and get instant AI-powered risk insights.",
        "Загрузите договор и получите мгновенные AI-подсказки по рискам.",
        "Келісімшартты жүктеп, тәуекелдер бойынша AI талдауын алыңыз.",
    ),
    "product": (
        "Upload your contract and get instant AI-powered risk insights.",
        "Загрузите договор и получите мгновенные AI-подсказки по рискам.",
        "Келісімшартты жүктеп, тәуекелдер бойынша AI талдауын алыңыз.",
    ),
    "about": (
        "Learn how TrustDoc.ai brings AI-powered legal clarity to every contract.",
        "Узнайте, как TrustDoc.ai дает юридическую ясность в каждом контракте.",
        "TrustDoc.ai әрбір келісімшартқа заңдық айқындық береді.",
    ),
    "project": (
        "See the TrustDoc.ai project vision, milestones, and next steps.",
        "Узнайте о видении проекта TrustDoc.ai, вехах и следующих шагах.",
        "TrustDoc.ai жобасының көзқарасы, кезеңдері және келесі қадамдары.",
    ),
    "news": (
        "Stay updated with TrustDoc.ai announcements and product releases.",
        "Следите за обновлениями и релизами продукта TrustDoc.ai.",
        "TrustDoc.ai жаңалықтары мен релиздері туралы біліңіз.",
    ),
    "gallery": (
        "Explore TrustDoc.ai visuals, previews, and interface highlights.",
        "Изучите визуалы TrustDoc.ai, превью и элементы интерфейса.",
        "TrustDoc.ai визуалдарын, алдын ала қарауларын зерттеңіз.",
    ),
    "contacts": (
        "Reach the TrustDoc.ai team for partnerships and support.",
        "Свяжитесь с командой TrustDoc.ai по партнерствам и поддержке.",
        "TrustDoc.ai командасымен серіктестік және қолдау үшін хабарласыңыз.",
    ),
    "freelancers": (
        "TrustDoc.ai for freelancers: protect projects and reduce legal friction.",
        "TrustDoc.ai для фрилансеров: защитите проекты и снижайте юридические риски.",
        "TrustDoc.ai фрилансерлер үшін: жобаларды қорғаңыз, құқықтық тәуекелдерді азайтыңыз.",
    ),
    "startups": (
        "TrustDoc.ai for startups: faster contracts, clearer risk insights.",
        "TrustDoc.ai для стартапов: быстрее заключайте контракты и видите риски.",
        "TrustDoc.ai стартаптар үшін: келісімдерді жеделдетіңіз, тәуекелдерді түсініңіз.",
    ),
    "news_detail": (
        "TrustDoc.ai update and announcement.",
        "Обновление и анонс TrustDoc.ai.",
        "TrustDoc.ai жаңалығы мен хабарламасы.",
    ),
}

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "dev-secret")

CLIENT_ID = os.getenv("THREADS_CLIENT_ID")
CLIENT_SECRET = os.getenv("THREADS_CLIENT_SECRET")
REDIRECT_URI = os.getenv("THREADS_REDIRECT_URI")

AUTH_URL = "https://www.threads.net/oauth/authorize"
TOKEN_URL = "https://graph.threads.net/oauth/access_token"
API_BASE = "https://graph.threads.net/v1.0"
CORS(app)

_openai_client: OpenAI | None = None


def _metadata(page: str) -> Dict[str, str]:
    title_en, title_ru, title_kz = TITLE_COPY[page]
    desc_en, desc_ru, desc_kz = DESCRIPTION_COPY[page]
    return {
        "title_en": title_en,
        "title_ru": title_ru,
        "title_kz": title_kz,
        "description_en": desc_en,
        "description_ru": desc_ru,
        "description_kz": desc_kz,
    }


def _get_openai_client() -> OpenAI:
    global _openai_client
    if _openai_client is None:
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise RuntimeError("OpenAI API key is not configured.")
        _openai_client = OpenAI(api_key=api_key)
    return _openai_client


def _extract_text(file_storage, extension: str) -> str:
    file_storage.stream.seek(0)
    data = file_storage.read()
    file_storage.stream.seek(0)

    if not data:
        raise ValueError("The uploaded file is empty.")

    if extension == ".pdf":
        reader = PyPDF2.PdfReader(BytesIO(data))
        text_chunks = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(text_chunks).strip()
        if not text:
            raise ValueError("No readable text found in the PDF document.")
        return text

    if extension == ".docx":
        document = docx.Document(BytesIO(data))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        table_cells = []
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        table_cells.append(cell_text)
        combined = "\n".join(paragraphs + table_cells).strip()
        if not combined:
            raise ValueError("No readable text found in the DOCX document.")
        return combined

    raise ValueError("Unsupported file type.")


def _analyze_with_openai(text: str) -> str:
    cleaned = text.strip()
    if not cleaned:
        raise ValueError("The document does not contain readable text.")

    snippet = cleaned if len(cleaned) <= 16000 else cleaned[:16000]

    try:
        detected_lang = langdetect.detect(snippet)
    except LangDetectException:
        detected_lang = "en"

    detected_lang = (detected_lang or "en").lower()
    if detected_lang not in {"en", "ru", "kk"}:
        detected_lang = "en"

    client = _get_openai_client()
    response = client.responses.create(
        model=OPENAI_MODEL,
        input=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {
                "role": "user",
                "content": f"The document language is {detected_lang}. Analyze accordingly:\n{snippet}",
            },
        ],
        temperature=0.2,
        max_output_tokens=900,
    )

    analysis = (response.output_text or "").strip()
    if not analysis:
        raise RuntimeError("Received an empty response from the AI service.")
    return analysis
    
@app.route("/threads-review")
def threads_review():
    return render_template("threads_review.html", results=None)

@app.route("/threads/login")
def threads_login():
    url = (
        f"{AUTH_URL}"
        f"?client_id={CLIENT_ID}"
        f"&redirect_uri={REDIRECT_URI}"
        f"&response_type=code"
        f"&scope=threads_basic,threads_keyword_search"
    )
    return redirect(url)

@app.route("/threads/callback")
def threads_callback():
    code = request.args.get("code")

    token_response = requests.post(
        TOKEN_URL,
        data={
            "client_id": CLIENT_ID,
            "client_secret": CLIENT_SECRET,
            "redirect_uri": REDIRECT_URI,
            "code": code,
            "grant_type": "authorization_code",
        },
    ).json()

    session["threads_token"] = token_response.get("access_token")
    return redirect("/threads-review")

@app.route("/threads/search", methods=["POST"])
def threads_search():
    keyword = request.form.get("keyword")
    token = session.get("threads_token")

    response = requests.get(
        f"{API_BASE}/keyword_search",
        params={
            "q": keyword,
            "search_type": "TOP",
            "fields": "id,text,username,permalink,timestamp",
            "access_token": token,
        },
    )

    results = response.json().get("data", [])
    return render_template("threads_review.html", results=results)


@app.route("/")
def index() -> str:
    return render_template("index.html", **_metadata("index"))


@app.route("/product")
def product() -> str:
    return render_template("product.html", **_metadata("product"))


@app.route("/analyze", methods=["POST"])
def analyze():
    file = request.files.get("file")
    if not file or file.filename == "":
        return jsonify({"status": "error", "message": "File is required."}), 400

    extension = os.path.splitext(file.filename)[1].lower()
    if extension not in ALLOWED_EXTENSIONS:
        return jsonify({"status": "error", "message": "Unsupported file type."}), 400

    try:
        text = _extract_text(file, extension)
    except ValueError as exc:
        return jsonify({"status": "error", "message": str(exc)}), 400
    except Exception as exc:  # pragma: no cover - defensive fallback
        return jsonify({"status": "error", "message": f"Failed to read document: {exc}"}), 500

    try:
        analysis = _analyze_with_openai(text)
    except ValueError as exc:
        return jsonify({"status": "error", "message": str(exc)}), 400
    except RuntimeError as exc:
        return jsonify({"status": "error", "message": str(exc)}), 500
    except Exception as exc:  # pragma: no cover - external dependency failure
        return jsonify({"status": "error", "message": f"AI analysis failed: {exc}"}), 502

    html_output = markdown.markdown(analysis, extensions=["extra", "tables"])

    return jsonify({"result": html_output})


if __name__ == "__main__":
    import os
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)

app.register_blueprint(threads_review_bp)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "dev-secret")
